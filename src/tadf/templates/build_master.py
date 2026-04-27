"""Build the docxtpl master templates programmatically — one per audit subtype.

The Phase 1 masters are generated from code so we can version-control them as a
script rather than as opaque binaries. Once the template stabilises and Fjodor
wants visual tweaks (fonts, headers, signatures with images) we switch to
hand-edited .docx per subtype — but for the MVP this keeps the loop fast.

Run via:
    uv run python -m tadf.templates.build_master
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TEMPLATES_DIR = Path(__file__).parent
SUBTYPES = ("kasutuseelne", "korraline", "erakorraline")


def _h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.size = Pt(14 if level == 1 else 12)


def _p(doc: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _add_field(paragraph, instr: str, font_size: int = 9) -> None:
    """Append a Word field (e.g. PAGE, NUMPAGES) to `paragraph`.

    python-docx has no high-level API for this, so we drop in the four
    `<w:fldChar>` / `<w:instrText>` elements that Word uses to recognise
    a field. We don't precompute the visible text — Word evaluates the
    field on open. LibreOffice (used by the PDF converter) supports the
    same encoding.
    """
    run = paragraph.add_run()
    run.font.size = Pt(font_size)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instr} "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    # Placeholder visible text — Word/LibreOffice replaces this on open.
    placeholder = OxmlElement("w:t")
    placeholder.text = "?"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _add_run(paragraph, text: str, *, font_size: int = 9, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold


def _setup_header_footer(doc: Document) -> None:
    """Attach a docxtpl-templated header / footer to the document's only section.

    Wires two placeholders that the context_builder fills:
      - `{{ page_header }}` — multi-line text (auditor override on «Готовый
        отчёт» if set, else `default_header_text(audit)` — Töö nr / Töö
        nimetus, mirroring the corpus convention).
      - `{{ page_footer }}` — single line with the päädev isik signatures.

    Plus a fixed page-number field (`Lk PAGE / NUMPAGES`) on a second header
    line, right-aligned. Word evaluates the field on open; LibreOffice does
    too (used for the PDF converter).

    The cover page (page 1) gets a DIFFERENT, empty header/footer via the
    `different_first_page_header_footer` flag, so the title page stays clean.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Default header (page 2+) — clear python-docx's stub paragraph first.
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Line 1: auditor-overridable text. The placeholder may contain newlines
    # (e.g. "Töö nr.: 1/2026\nTöö nimetus: …") which docxtpl renders as soft
    # line breaks inside this single paragraph.
    h1 = header.add_paragraph()
    _add_run(h1, "{{ page_header }}")

    # Line 2: page-of-pages on the right via a right-aligned tab stop.
    h2 = header.add_paragraph()
    h2.paragraph_format.tab_stops.add_tab_stop(
        Cm(16), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    _add_run(h2, "\tLk ")
    _add_field(h2, "PAGE")
    _add_run(h2, " / ")
    _add_field(h2, "NUMPAGES")

    # Default footer (page 2+).
    footer = section.footer
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    f1 = footer.add_paragraph()
    _add_run(f1, "{{ page_footer }}")

    # First-page header/footer — empty paragraph so the cover stays clean.
    # python-docx requires the parts to exist before we touch them, so we
    # explicitly enable the slot and leave it default-empty.
    first_header = section.first_page_header
    for p in list(first_header.paragraphs):
        p._element.getparent().remove(p._element)
    first_header.add_paragraph()  # empty

    first_footer = section.first_page_footer
    for p in list(first_footer.paragraphs):
        p._element.getparent().remove(p._element)
    first_footer.add_paragraph()  # empty


def build(subtype: str = "kasutuseelne") -> Path:
    doc = Document()

    # Header + footer FIRST so they're attached to section 0 before any body
    # content is added (python-docx works either order, but doing it first
    # makes the docxtpl render order obvious to the next reader).
    _setup_header_footer(doc)

    # ---------- 0. TITTELLEHT ----------
    _p(doc, "{{ cover.title }}", bold=True, center=True)
    _p(doc, "Aadress: {{ building.address }}", center=True)
    _p(doc, "Ehitisregistri kood: {{ building.ehr_code }}", center=True)
    _p(doc, "Töö nr: {{ audit.display_no }}", center=True)
    doc.add_paragraph()
    _p(
        doc,
        "Auditi koostas: {{ composer.full_name }}, "
        "{% if composer.company %}{{ composer.company }}"
        "{% if composer.company_reg_nr %}, reg. nr {{ composer.company_reg_nr }}{% endif %}"
        "{% endif %}",
    )
    _p(
        doc,
        "Auditi kontrollis (vastutav pädev isik): {{ reviewer.full_name }}, "
        "kutsetunnistus {{ reviewer.kutsetunnistus_no }}"
        "{% if reviewer.qualification %}, {{ reviewer.qualification }}{% endif %}",
    )
    _p(doc, "Tellija: {{ client.name if client else '' }}")
    doc.add_paragraph()
    _p(doc, "{{ visit_date_str }}", center=True)

    doc.add_page_break()

    # ---------- 1. ÜLDOSA ----------
    _h(doc, "1. ÜLDOSA")

    _h(doc, "1.1. Ehitise auditi eesmärk", level=2)
    _p(doc, "{{ audit.purpose }}")

    _h(doc, "1.2. Ehitise auditori andmed", level=2)
    _p(doc, "Auditi liik: {{ audit_type_text }}")
    _p(doc, "Auditi koostas: {{ composer.full_name }}")
    _p(
        doc,
        "Vastutav pädev isik: {{ reviewer.full_name }}, kutsetunnistus {{ reviewer.kutsetunnistus_no }}",
    )
    _p(doc, "{{ independence_declaration }}")

    _h(doc, "1.3. Auditi ulatus", level=2)
    _p(doc, "{{ audit.scope }}")

    # ---------- 2. AUDITI OBJEKT ----------
    _h(doc, "2. AUDITI OBJEKT JA SELLE KIRJELDUS")
    _p(doc, "Objekti aadress: {{ building.address }}")
    _p(doc, "Ehitisregistri kood: {{ building.ehr_code }}")
    _p(doc, "Kasutusotstarve: {{ building.use_purpose }}")
    _p(doc, "Ehitusaasta: {{ building.construction_year }}")

    # ---------- 3. KINNISTU ----------
    _h(doc, "3. KINNISTU ASUKOHT JA PLANEERING")
    _p(doc, "Katastritunnus: {{ building.kataster_no }}")
    _p(doc, "Kinnistu pindala: {{ building.site_area_m2 }} m²")

    # ---------- 4. HOONE ÜLEVAATUS ----------
    _h(doc, "4. HOONE ÜLEVAATUS")
    _p(doc, "Paikvaatlus toimus {{ visit_date_str }} visuaalkontrolli teel.")
    doc.add_paragraph("{% for f in findings_section_4 %}")
    _p(doc, "• {{ f.observation }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 5-6. ARHITEKTUUR + KONSTRUKTSIOON ----------
    _h(doc, "5. HOONE ARHITEKTUURI- JA EHITUSLIK OSA")
    doc.add_paragraph("{% for f in findings_section_5 %}")
    _p(doc, "• {{ f.observation }}")
    doc.add_paragraph("{% endfor %}")

    _h(doc, "6. HOONE KONSTRUKTIIVNE OSA")
    doc.add_paragraph("{% for f in findings_section_6 %}")
    _p(doc, "• {{ f.observation }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 7. TEHNOSÜSTEEMID ----------
    _h(doc, "7. HOONE TEHNOSÜSTEEMID")
    doc.add_paragraph("{% for f in findings_section_7 %}")
    _p(doc, "• {{ f.observation }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 8. TULEKAITSE (conditional) ----------
    doc.add_paragraph("{% if building.fire_class %}")
    _h(doc, "8. HOONE TULEKAITSE OSA")
    _p(doc, "Tulepüsivusklass: {{ building.fire_class }}")
    doc.add_paragraph("{% for f in findings_section_8 %}")
    _p(doc, "• {{ f.observation }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endif %}")

    # ---------- 10. TEHNILISED NÄITAJAD ----------
    _h(doc, "10. HOONE TEHNILISED NÄITAJAD")
    _p(doc, "Ehitisealune pind: {{ building.footprint_m2 }} m²")
    _p(doc, "Maht: {{ building.volume_m3 }} m³")
    _p(doc, "Maapealsete korruste arv: {{ building.storeys_above }}")
    _p(doc, "Maa-aluste korruste arv: {{ building.storeys_below }}")

    # ---------- 11. KOKKUVÕTE (auditor-only) ----------
    _h(doc, "11. KOKKUVÕTE")
    doc.add_paragraph("{% for f in findings_section_11 %}")
    _p(doc, "{{ f.observation }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 12. ÕIGUSLIKUD ALUSED ----------
    _h(doc, "12. AUDITI ÕIGUSLIKUD ALUSED JA ULATUS")
    _p(doc, "Käesolev audit on koostatud tuginedes:")
    doc.add_paragraph("{% for r in legal_refs %}")
    _p(doc, "• {{ r.code }} — {{ r.title_et }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 13. METOODIKA ----------
    _h(doc, "13. AUDITI METOODIKA")
    _p(doc, "{{ methodology }}")

    # ---------- 14. LÕPPHINNANG (auditor-only) ----------
    _h(doc, "14. AUDITI LÕPPHINNANG")
    doc.add_paragraph("{% for f in findings_section_14 %}")
    _p(doc, "{{ f.observation }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 16. FOTOD ----------
    _h(doc, "16. FOTOD")
    doc.add_paragraph("{% for photo in photos %}")
    _p(doc, "[{{ photo.section_ref }}]  {{ photo.caption }}")
    doc.add_paragraph("{{ photo.image }}")
    doc.add_paragraph("{% endfor %}")

    # ---------- 15. ALLKIRJAD ----------
    _h(doc, "15. ALLKIRJAD")
    _p(doc, "Auditi koostas:")
    _p(
        doc,
        "{{ composer.full_name }}"
        "{% if composer.company %} ({{ composer.company }}"
        "{% if composer.company_reg_nr %}, reg. nr {{ composer.company_reg_nr }}{% endif %})"
        "{% endif %}",
    )
    _p(doc, "allkiri: digitaalselt")
    doc.add_paragraph()
    _p(doc, "Auditi kontrollis (vastutav pädev isik):")
    _p(doc, "{{ reviewer.full_name }}, kutsetunnistus {{ reviewer.kutsetunnistus_no }}")
    _p(doc, "allkiri: digitaalselt")
    doc.add_paragraph()
    _p(doc, "{{ retention_notice }}")

    out = TEMPLATES_DIR / f"ea_{subtype}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_all() -> list[Path]:
    """Build one template per subtype. Bodies are identical today; the variation
    lives in `boilerplate.yaml` (audit_purpose) and `context_builder` (label)."""
    return [build(s) for s in SUBTYPES]


if __name__ == "__main__":
    for p in build_all():
        print(f"wrote {p}")
