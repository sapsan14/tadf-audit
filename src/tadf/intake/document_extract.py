"""Convert an uploaded project document (DOCX/PDF/DOC/ASiC-E) to plain text
and run it through the Building extractor.

Public entry point: `extract_from_upload(uploaded_file) -> (extracted_dict, raw_text)`.
The raw_text is returned alongside the extraction so the UI can show
Fjodor what Claude actually saw (debug expander).

Supported formats:
  - `.docx` — direct python-docx parse
  - `.pdf`  — direct pdfplumber parse
  - `.doc`  — converted to docx via headless LibreOffice (`soffice`),
    then routed back to the docx path. Raises `LibreofficeMissing` if
    LibreOffice is not on PATH.
  - `.asice` — unzipped, inner .docx/.pdf extracted, routed by inner type.
"""

from __future__ import annotations

import tempfile
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any, Protocol

from tadf.corpus.parse_doc import LibreofficeMissing, _find_soffice
from tadf.llm.extractor import extract_building

# Re-export so callers can `except LibreofficeMissing` without reaching
# into the corpus module.
__all__ = ["LibreofficeMissing", "extract_from_upload", "to_text"]


class _UploadedFile(Protocol):
    """Subset of streamlit.UploadedFile / file-like."""

    name: str

    def read(self) -> bytes: ...


def _docx_to_text(data: bytes) -> str:
    """Concatenate every paragraph in a .docx file. Project explanatory
    notes don't follow the §5 audit outline, so we skip the heading-aware
    heuristics in `corpus/parse_docx.py` and just take everything."""
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # Tables often hold the most structured data (year, area, etc.) in
    # explanatory notes — flatten them.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _pdf_to_text(data: bytes) -> str:
    """Concatenate every page's text in a .pdf via pdfplumber."""
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            t = t.strip()
            if t:
                parts.append(t)
    return "\n".join(parts)


def _doc_to_text(data: bytes) -> str:
    """Convert legacy .doc bytes to plain text via headless LibreOffice.

    Writes the bytes to a temp file, runs `soffice --convert-to docx`, then
    delegates to `_docx_to_text`. Raises `LibreofficeMissing` cleanly when
    soffice is absent so callers can show a tailored UI message.
    """
    import subprocess

    soffice = _find_soffice()
    if soffice is None:
        raise LibreofficeMissing(
            "soffice/libreoffice not found on PATH. Install LibreOffice "
            "(apt-get install libreoffice) to import .doc files."
        )

    with tempfile.TemporaryDirectory(prefix="tadf-intake-doc-") as tmp_dir:
        tmp_path = Path(tmp_dir)
        src = tmp_path / "input.doc"
        src.write_bytes(data)
        proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", str(tmp_path), str(src)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode != 0:
            raise RuntimeError(
                f"LibreOffice failed to convert .doc: rc={proc.returncode} "
                f"stderr={proc.stderr.strip()[:200]}"
            )
        converted = tmp_path / "input.docx"
        if not converted.exists():
            raise RuntimeError(
                "LibreOffice reported success but the converted .docx is missing"
            )
        return _docx_to_text(converted.read_bytes())


def _asice_to_text(data: bytes) -> str:
    """Unzip the ASiC-E container, find the first .docx/.pdf payload, and
    convert it. Skips signature XMLs, the mimetype marker, and META-INF."""
    with zipfile.ZipFile(BytesIO(data)) as z:
        names = [
            n
            for n in z.namelist()
            if not n.startswith("META-INF/")
            and not n.lower().endswith(".xml")
            and n != "mimetype"
        ]
        chosen = next((n for n in names if n.lower().endswith(".docx")), None)
        if chosen is None:
            chosen = next((n for n in names if n.lower().endswith(".pdf")), None)
        if chosen is None:
            chosen = next((n for n in names if n.lower().endswith(".doc")), None)
        if chosen is None:
            raise ValueError(
                "No .docx / .pdf / .doc payload found inside the ASiC-E container."
            )
        inner = z.read(chosen)

    inner_ext = Path(chosen).suffix.lower()
    if inner_ext == ".docx":
        return _docx_to_text(inner)
    if inner_ext == ".pdf":
        return _pdf_to_text(inner)
    if inner_ext == ".doc":
        return _doc_to_text(inner)
    raise ValueError(f"Unsupported inner type in ASiC-E: {inner_ext}")


def to_text(name: str, data: bytes) -> str:
    """Detect file type by extension, return plain text. Raises ValueError
    on unsupported extensions; raises `LibreofficeMissing` for `.doc` (and
    for `.asice` whose payload is `.doc`) when soffice isn't available."""
    ext = Path(name).suffix.lower()
    if ext == ".docx":
        return _docx_to_text(data)
    if ext == ".pdf":
        return _pdf_to_text(data)
    if ext == ".doc":
        return _doc_to_text(data)
    if ext == ".asice":
        return _asice_to_text(data)
    raise ValueError(
        f"Unsupported file type: {ext} (supported: .docx, .pdf, .doc, .asice)"
    )


def extract_from_upload(uploaded_file: _UploadedFile) -> tuple[dict[str, Any], str]:
    """Read the file, convert to text, run extractor. Returns
    (extracted_dict, raw_text)."""
    data = uploaded_file.read()
    text = to_text(uploaded_file.name, data)
    extracted = extract_building(text)
    return extracted, text
