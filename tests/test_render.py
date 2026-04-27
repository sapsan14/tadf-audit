from __future__ import annotations

import json

import pytest
from docx import Document

from tadf.render.docx_render import ChecklistFailed, render_to_path


def test_render_writes_docx_and_context(audit, tmp_path):
    out = render_to_path(audit, tmp_path)
    assert out.exists()
    assert out.stat().st_size > 5_000  # not an empty stub
    assert (tmp_path / "context.json").exists()
    ctx = json.loads((tmp_path / "context.json").read_text(encoding="utf-8"))
    assert ctx["composer"]["full_name"] == "Aleksei Sholokhov"
    assert ctx["reviewer"]["kutsetunnistus_no"] == "148515"


def test_render_contains_key_estonian_strings(audit, tmp_path):
    out = render_to_path(audit, tmp_path)
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)

    # Section headings
    assert "ÜLDOSA" in body
    assert "AUDITI OBJEKT JA SELLE KIRJELDUS" in body
    assert "TULEKAITSE" in body  # conditional on fire_class
    assert "KOKKUVÕTE" in body
    assert "LÕPPHINNANG" in body

    # Filled values
    assert "148515" in body
    assert "Fjodor Sokolov" in body
    assert "Aleksei Sholokhov" in body
    assert "Linna AÜ 1062" in body
    assert "TP-3" in body

    # Boilerplate
    assert "kutsetunnistus" in body.lower()
    assert "sõltumatud" in body  # independence_declaration
    assert "metoodili" in body.lower()  # methodology block


def test_render_blocks_when_checklist_fails(audit, tmp_path):
    audit.purpose = None
    with pytest.raises(ChecklistFailed):
        render_to_path(audit, tmp_path, enforce_checklist=True)


def test_render_force_when_checklist_disabled(audit, tmp_path):
    audit.purpose = None
    out = render_to_path(audit, tmp_path, enforce_checklist=False)
    assert out.exists()


def test_no_fire_class_skips_section_8(audit, tmp_path):
    audit.building.fire_class = None
    audit.findings = [f for f in audit.findings if not f.section_ref.startswith("8")]
    out = render_to_path(audit, tmp_path)
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "TULEKAITSE" not in body


def _header_footer_texts(docx_path) -> tuple[str, str]:
    """Return ("header text on page 2+", "footer text on page 2+") by reading
    whichever header/footer part `sectPr` tags as `type="default"`."""
    import re
    import zipfile

    with zipfile.ZipFile(str(docx_path)) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        rel_to_target = {
            m.group("id"): m.group("target")
            for m in re.finditer(
                r'<Relationship\b[^>]*?\bId="(?P<id>[^"]+)"[^>]*?\bTarget="(?P<target>[^"]+)"',
                rels,
            )
        }

        sect = re.search(r"<w:sectPr.*?</w:sectPr>", z.read("word/document.xml").decode("utf-8"), re.DOTALL)
        assert sect is not None
        h_default = re.search(r'<w:headerReference w:type="default" r:id="([^"]+)"', sect.group(0))
        f_default = re.search(r'<w:footerReference w:type="default" r:id="([^"]+)"', sect.group(0))
        assert h_default and f_default, "default header/footer references missing"

        def _text(part_path: str) -> str:
            xml = z.read(f"word/{part_path}").decode("utf-8")
            return " ".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))

        return _text(rel_to_target[h_default.group(1)]), _text(rel_to_target[f_default.group(1)])


def test_render_populates_header_with_audit_metadata(audit, tmp_path):
    """The header on page 2+ must show the page_header text computed by
    context_builder.default_header_text — Töö nr + Töö nimetus + address —
    plus a Lk PAGE/NUMPAGES field."""
    out = render_to_path(audit, tmp_path)
    header_text, _ = _header_footer_texts(out)
    assert "Töö nr" in header_text
    assert "Töö nimetus" in header_text
    assert "Linna AÜ 1062" in header_text  # building.address
    # The PAGE / NUMPAGES Word fields render the literal "?" we drop in as
    # a placeholder; Word/LibreOffice evaluates them on open.
    assert "Lk" in header_text


def test_render_populates_footer_with_signatures(audit, tmp_path):
    """Footer on page 2+ must show the päädev isik signature line computed
    by context_builder.default_footer_text."""
    out = render_to_path(audit, tmp_path)
    _, footer_text = _header_footer_texts(out)
    assert "Pädev isik" in footer_text
    assert "Fjodor Sokolov" in footer_text
    assert "148515" in footer_text  # kutsetunnistus number


def test_header_footer_use_auditor_override_when_set(audit, tmp_path):
    """When the auditor types a custom header/footer in the «Готовый отчёт»
    page, that text wins over the computed default."""
    audit.header_override = "CUSTOM HEADER OVERRIDE TEXT"
    audit.footer_override = "CUSTOM FOOTER OVERRIDE TEXT"
    out = render_to_path(audit, tmp_path)
    header_text, footer_text = _header_footer_texts(out)
    assert "CUSTOM HEADER OVERRIDE TEXT" in header_text
    assert "CUSTOM FOOTER OVERRIDE TEXT" in footer_text
    # The computed default values should NOT have been rendered.
    assert "Pädev isik" not in footer_text


def test_render_first_page_header_is_empty(audit, tmp_path):
    """The cover page must NOT inherit the page-2+ header — Word's
    `titlePg` flag plus an empty `first_page_header` keep the title page
    visually clean. Regress on this and Fjodor sees his work-nr stamped
    over the title.
    """
    import re
    import zipfile

    out = render_to_path(audit, tmp_path)
    with zipfile.ZipFile(str(out)) as z:
        sect = re.search(r"<w:sectPr.*?</w:sectPr>", z.read("word/document.xml").decode("utf-8"), re.DOTALL)
        assert sect is not None
        assert "<w:titlePg/>" in sect.group(0)
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        rel_to_target = {
            m.group("id"): m.group("target")
            for m in re.finditer(
                r'<Relationship\b[^>]*?\bId="(?P<id>[^"]+)"[^>]*?\bTarget="(?P<target>[^"]+)"',
                rels,
            )
        }
        h_first = re.search(r'<w:headerReference w:type="first" r:id="([^"]+)"', sect.group(0))
        assert h_first is not None
        first_xml = z.read(f"word/{rel_to_target[h_first.group(1)]}").decode("utf-8")
        first_text = " ".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", first_xml))
        assert first_text.strip() == ""
